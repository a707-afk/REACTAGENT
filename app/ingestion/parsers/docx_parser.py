"""DOCX parser using python-docx for text and table extraction."""
from __future__ import annotations

import logging
from typing import Any

from app.ingestion.parsers.base import BaseParser, ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parse DOCX files using python-docx."""

    supported_mimes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def parse(self, file_path: str, **kwargs) -> ParsedDocument:
        """Parse a DOCX file."""
        import docx

        doc = docx.Document(file_path)
        paragraphs = []
        tables_md = []
        page_number = 1  # DOCX doesn't have real pages, treat as single page

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading levels
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading 1" in style_name:
                    paragraphs.append(f"# {text}")
                elif "heading 2" in style_name:
                    paragraphs.append(f"## {text}")
                elif "heading 3" in style_name:
                    paragraphs.append(f"### {text}")
                elif "heading 4" in style_name:
                    paragraphs.append(f"#### {text}")
                else:
                    paragraphs.append(text)

        # Extract tables
        for table in doc.tables:
            md_table = self._table_to_markdown(table)
            if md_table:
                tables_md.append(md_table)

        # Extract images (with VLM OCR if available)
        images = []
        try:
            images = self._extract_images_with_ocr(doc)
        except Exception as e:
            logger.warning("Image extraction failed: %s", e)

        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            pages=[
                ParsedPage(
                    page_number=1,
                    text=full_text,
                    tables=tables_md,
                    images=images,
                    metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables)},
                )
            ],
            total_pages=1,
            metadata={"file_path": file_path},
        )

    def _table_to_markdown(self, table) -> str:
        """Convert a python-docx table to Markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        max_cols = max(len(row) for row in rows)
        for row in rows:
            while len(row) < max_cols:
                row.append("")

        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _extract_images_with_ocr(self, doc) -> list[dict[str, Any]]:
        """Extract images from DOCX and OCR them using VLM."""
        import io
        from zipfile import ZipFile

        images = []
        try:
            from app.vlm import ocr_image
            has_vlm = True
        except ImportError:
            has_vlm = False
            logger.info("VLM not available, skipping image OCR in DOCX")

        # DOCX is a ZIP file; images are in word/media/
        try:
            with ZipFile(doc.part.blob if hasattr(doc.part, 'blob') else io.BytesIO()) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        img_bytes = zf.read(name)
                        ocr_text = ""
                        if has_vlm:
                            try:
                                ocr_text = ocr_image(image_bytes=img_bytes)
                            except Exception as e:
                                logger.warning("VLM OCR failed for %s: %s", name, e)
                        images.append({
                            "page": 1,
                            "bbox": [0, 0, 0, 0],
                            "ocr_text": ocr_text,
                            "source": name,
                        })
        except Exception as e:
            logger.debug("Could not extract images from DOCX: %s", e)

        return images
